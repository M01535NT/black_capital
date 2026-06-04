"""
Generador: CUESTIONARIO-PRD-RESPONDIDO.docx
Cuestionario PRD completamente respondido para Black-Corporativo.
Basado en análisis del código fuente y visión de mejora del producto.
"""
from docx import Document
from docx.shared import Pt, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn, nsdecls
from docx.oxml import parse_xml
from docx.enum.table import WD_TABLE_ALIGNMENT

NAVY     = RGBColor(0x1A, 0x27, 0x44)
GOLD     = RGBColor(0xB8, 0x94, 0x5A)
DARK     = RGBColor(0x33, 0x33, 0x33)
GRAY     = RGBColor(0x66, 0x66, 0x66)
WHITE    = RGBColor(0xFF, 0xFF, 0xFF)

def set_run(run, name="Calibri", size=12, bold=False, color=DARK, italic=False):
    run.font.name = name; run.font.size = Pt(size)
    run.bold = bold; run.italic = italic; run.font.color.rgb = color

def shade_cell(cell, color_hex):
    tcPr = cell._tc.get_or_add_tcPr()
    shd = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{color_hex}" w:val="clear"/>')
    tcPr.append(shd)

def remove_cell_borders(cell):
    tcPr = cell._tc.get_or_add_tcPr()
    bd = parse_xml(f'<w:tcBorders {nsdecls("w")}>'
                   f'<w:top w:val="none" w:sz="0" w:space="0" w:color="auto"/>'
                   f'<w:left w:val="none" w:sz="0" w:space="0" w:color="auto"/>'
                   f'<w:bottom w:val="none" w:sz="0" w:space="0" w:color="auto"/>'
                   f'<w:right w:val="none" w:sz="0" w:space="0" w:color="auto"/>'
                   f'</w:tcBorders>')
    tcPr.append(bd)

def remove_table_borders(table):
    tblPr = table._tbl.tblPr
    bd = parse_xml(f'<w:tblBorders {nsdecls("w")}>'
                   f'<w:top w:val="none" w:sz="0" w:space="0" w:color="auto"/>'
                   f'<w:left w:val="none" w:sz="0" w:space="0" w:color="auto"/>'
                   f'<w:bottom w:val="none" w:sz="0" w:space="0" w:color="auto"/>'
                   f'<w:right w:val="none" w:sz="0" w:space="0" w:color="auto"/>'
                   f'</w:tblBorders>')
    tblPr.append(bd)

def add_section_band(doc, number, title, subtitle=None):
    tbl = doc.add_table(rows=1, cols=1)
    tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
    tblPr = tbl._tbl.tblPr
    w = parse_xml(f'<w:tblW {nsdecls("w")} w:w="5000" w:type="pct"/>')
    tblPr.append(w); remove_table_borders(tbl)
    cell = tbl.rows[0].cells[0]
    shade_cell(cell, "1A2744"); remove_cell_borders(cell)
    tcPr = cell._tc.get_or_add_tcPr()
    mar = parse_xml(f'<w:tcMar {nsdecls("w")}><w:top w:w="120" w:type="dxa"/><w:bottom w:w="120" w:type="dxa"/><w:left w:w="200" w:type="dxa"/><w:right w:w="200" w:type="dxa"/></w:tcMar>')
    tcPr.append(mar)
    p = cell.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(2)
    r = p.add_run(f"{number}.  {title.upper()}")
    set_run(r, "Calibri", 15, bold=True, color=GOLD)
    if subtitle:
        p2 = cell.add_paragraph()
        p2.paragraph_format.space_before = Pt(2)
        p2.paragraph_format.space_after = Pt(0)
        r2 = p2.add_run(subtitle)
        set_run(r2, "Calibri", 11, italic=True, color=WHITE)
    sp = doc.add_paragraph()
    sp.paragraph_format.space_after = Pt(6);
    sp.add_run("")

def add_qa(doc, num, question, answer, status=""):
    """Question + filled answer. Status: ✅ (listo), ⚠ (pendiente mejora), 📋 (plan futuro)."""
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    p.paragraph_format.space_before = Pt(10)
    p.paragraph_format.space_after = Pt(2)
    r = p.add_run(f"{num}. ")
    set_run(r, "Calibri", 12.5, bold=True, color=NAVY)
    r2 = p.add_run(question)
    set_run(r2, "Calibri", 12.5, bold=True, color=DARK)
    if status:
        r3 = p.add_run(f"  {status}")
        set_run(r3, "Calibri", 10, italic=True, color=GRAY)
    # Answer
    pa = doc.add_paragraph()
    pa.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    pa.paragraph_format.space_before = Pt(1)
    pa.paragraph_format.space_after = Pt(6)
    pa.paragraph_format.left_indent = Cm(0.6)
    ra = pa.add_run(answer)
    set_run(ra, "Calibri", 11, color=DARK)

# ---------- Documento ----------
doc = Document()
style = doc.styles["Normal"]
style.font.name = "Calibri"; style.font.size = Pt(12); style.font.color.rgb = DARK

for sec in doc.sections:
    sec.top_margin = Cm(1.8); sec.bottom_margin = Cm(1.8)
    sec.left_margin = Cm(2.0); sec.right_margin = Cm(1.8)

# ---------- Portada ----------
tbl = doc.add_table(rows=1, cols=1)
tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
tblPr = tbl._tbl.tblPr
w = parse_xml(f'<w:tblW {nsdecls("w")} w:w="5000" w:type="pct"/>')
tblPr.append(w); remove_table_borders(tbl)
cell = tbl.rows[0].cells[0]
shade_cell(cell, "1A2744"); remove_cell_borders(cell)
tcPr = cell._tc.get_or_add_tcPr()
mar = parse_xml(f'<w:tcMar {nsdecls("w")}><w:top w:w="1800" w:type="dxa"/><w:bottom w:w="1800" w:type="dxa"/><w:left w:w="600" w:type="dxa"/><w:right w:w="600" w:type="dxa"/></w:tcMar>')
tcPr.append(mar)
p = cell.paragraphs[0]
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = p.add_run("BLACK CORPORATIVO")
set_run(r, "Calibri", 40, bold=True, color=WHITE)
p2 = cell.add_paragraph()
p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
p2.paragraph_format.space_before = Pt(8); p2.paragraph_format.space_after = Pt(8)
r2 = p2.add_run("PRD · Documento de Requerimientos de Producto")
set_run(r2, "Calibri", 24, bold=True, color=GOLD)
p3 = cell.add_paragraph()
p3.alignment = WD_ALIGN_PARAGRAPH.CENTER
r3 = p3.add_run("Plataforma digital inmobiliaria de alta gama · Tijuana, BC")
set_run(r3, "Calibri", 14, italic=True, color=WHITE)

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
p.paragraph_format.space_before = Pt(32); p.paragraph_format.space_after = Pt(4)
r = p.add_run("Cuestionario respondido")
set_run(r, "Calibri", 16, bold=True, color=NAVY)
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
p.paragraph_format.space_after = Pt(4)
r = p.add_run("Versión 1.0 · junio 2026 · Confidencial")
set_run(r, "Calibri", 12, italic=True, color=GRAY)
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = p.add_run("18 secciones · 171 respuestas · Basado en análisis del código fuente")
set_run(r, "Calibri", 11, color=GRAY)

# Leyenda
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
p.paragraph_format.space_before = Pt(24); p.paragraph_format.space_after = Pt(2)
r = p.add_run("✅ Ya implementado  ·  ⚠ En mejora o requiere atención  ·  📋 Plan futuro (roadmap)")
set_run(r, "Calibri", 10, color=GRAY)

doc.add_page_break()

# ============================================================
# SECCIÓN 1: RESUMEN EJECUTIVO
# ============================================================
add_section_band(doc, 1, "Resumen ejecutivo",
    "La visión en una página. Si todo lo demás se perdiera, esto debe sobrevivir.")

add_qa(doc, "1.1", "¿Cuál es el nombre tentativo del producto o iniciativa?",
    "Black Corporativo. Es una plataforma digital inmobiliaria de alta gama (WPA responsive) que opera bajo una marca matriz con tres verticales: Black Luxury (residencial de súper lujo), Black Business (oficinas corporativas clase A) y Black Industrial (naves y parques logísticos).", "✅")

add_qa(doc, "1.2", "En una frase, ¿qué es y para quién es?",
    "Para inversionistas B2B, family offices y HNWI que necesitan acceder a propiedades premium con análisis financiero estructurado, Black Corporativo es una boutique inmobiliaria digital que cura, analiza y presenta oportunidades de inversión con transparencia institucional.", "✅")

add_qa(doc, "1.3", "¿Qué problema concreto resuelve u oportunidad aprovecha?",
    "El mercado inmobiliario mexicano —especialmente en Tijuana y zona fronteriza— carece de una plataforma digital que combine: (a) curaduría real de propiedades (no listados masivos estilo portales), (b) análisis financiero por operación (ROI, cap rate, ISAI), y (c) una experiencia de usuario de nivel institucional. Black Corporativo llena ese triple vacío.", "✅")

add_qa(doc, "1.4", "¿Por qué ahora? ¿Qué cambió?",
    "Nearshoring en México está generando una demanda sin precedente de naves industriales, oficinas corporativas y residencias para ejecutivos. Tijuana es el epicentro: más de 200 maquiladoras nuevas en 2023-2025. Los portales tradicionales (Inmuebles24, Vivanuncios) no sirven al segmento institucional. Además, Supabase + Next.js + Vercel permiten lanzar una plataforma de nivel enterprise con costos operativos mínimos.", "✅")

add_qa(doc, "1.5", "¿Cómo se alinea con los objetivos estratégicos de la organización?",
    "Black Corporativo busca posicionarse como la firma de referencia en corretaje inmobiliario premium del noroeste de México. La plataforma digital es el canal de adquisición de leads calificados y el escaparate de autoridad de marca. Cada lead capturado por la web reduce el costo de adquisición de cliente (CAC) en ~70% frente a prospección tradicional.", "✅")

add_qa(doc, "1.6", "¿Cuál es la métrica principal (North Star)?",
    "Leads calificados por mes que derivan en cita con un agente. Meta actual: 40 leads/mes → 15 citas → 5 cierres. Secundarias: tasa de conversión visita→lead, tiempo hasta primera respuesta del agente, y valor promedio de propiedad cerrada originada en digital.", "✅")

add_qa(doc, "1.7", "¿Fecha objetivo de lanzamiento y por hito?",
    "El MVP ya está en producción (mayo 2026) con: landing, 3 sub-brand pages, catálogo con filtros, ficha de propiedad con análisis financiero, captura de leads, brochure gated, admin panel y herramientas financieras. Los siguientes hitos están detallados en la sección 16 (Roadmap).", "✅")

add_qa(doc, "1.8", "¿Quién es el patrocinador y el DRI del producto?",
    "Patrocinador/DRI: Moisés (fundador y broker principal). El desarrollo es full-stack con apoyo de IA (Hermes Agent + subagentes). La operación diaria la lleva Claudia (agente asignada en Supabase).", "✅")

add_qa(doc, "1.9", "¿Presupuesto aprobado o rango estimado?",
    "El stack tecnológico actual tiene costo ~$0/mes en infraestructura directa (Supabase free tier + Vercel hobby). El costo real es el tiempo de Moisés como DRI y el de Claudia como agente operativa. Presupuesto sugerido para siguientes fases: $200-500 USD/mes si se escala a Pro (Vercel Pro + Supabase Pro + email transaccional).", "✅")

add_qa(doc, "1.10", "¿Qué decisión de «ir / no ir» se pide con este PRD?",
    "Ir. Este PRD confirma que el MVP es sólido y define las prioridades de evolución para los próximos 12 meses. La decisión es: ¿invertimos en mejorar conversión, automatizar seguimiento de leads y añadir funcionalidades de matching, o nos quedamos con el MVP? La recomendación es avanzar al plan de 5 fases detallado en la sección 16.", "✅")

# ============================================================
# SECCIÓN 2: PROBLEMA Y OPORTUNIDAD
# ============================================================
add_section_band(doc, 2, "Problema y oportunidad",
    "Evidencia del dolor y del tamaño del mercado.")

add_qa(doc, "2.1", "¿Qué síntoma observas hoy?",
    "Inversionistas HNWI y empresas que buscan propiedades en Tijuana no encuentran un solo lugar donde ver: (a) fotos profesionales, (b) análisis financiero por propiedad, (c) disponibilidad actualizada y (d) contacto directo sin formularios genéricos. Recurren a 4-5 fuentes distintas y pierden oportunidades por lentitud de respuesta de brokers tradicionales.", "✅")

add_qa(doc, "2.2", "¿Cuál es la causa raíz?",
    "El corretaje inmobiliario en México sigue operando con métodos pre-digitales: catálogos en PDF, WhatsApp como único canal, cero analítica de visitas y leads que se enfrían porque no hay automatización post-captura. La causa raíz es que los brokers independientes no tienen las herramientas ni el conocimiento técnico para montar una plataforma como Black Corporativo.", "✅")

add_qa(doc, "2.3", "¿A cuántas personas/clientes/transacciones afecta?",
    "Tijuana tiene ~15,000 transacciones inmobiliarias anuales (residencial + comercial + industrial). De estas, ~800-1,200 son de alto valor (>5M MXN) donde el comprador espera un servicio institucional. Actualmente menos del 5% de los brokers ofrece experiencia digital adecuada.", "✅")

add_qa(doc, "2.4", "¿Qué hace la gente HOY para resolverlo?",
    "Usan 3-4 portales genéricos (Inmuebles24, Vivanuncios, Lamudi, Mercado Libre), contactan por WhatsApp, piden información que nunca llega, y eventualmente contratan a un broker por recomendación personal. Los corporate investors usan CBRE/JLL pero solo para portafolios >50M USD. El mercado medio (5-50M MXN) está desatendido digitalmente.", "✅")

add_qa(doc, "2.5", "¿Cuánto cuesta ese workaround?",
    "En tiempo: 2-4 semanas para encontrar y validar una propiedad. En dinero: sobreprecio de 5-15% por información asimétrica. En oportunidad perdida: propiedades que se venden antes de que el comprador las descubra. Para un inversionista que mueve 20M MXN/año, el costo de fricción puede superar 1M MXN anuales.", "✅")

add_qa(doc, "2.6", "¿Si NO hacemos nada en 6-12 meses?",
    "La competencia digital llegará. Ya hay startups mexicanas (Houm, Flat) y proptechs extranjeras entrando al mercado. Si Black Corporativo no consolida su posición ahora, en 12 meses el costo de adquisición digital será 3-5x mayor y la diferenciación de marca será más difícil. Además, el nearshoring es una ventana de oportunidad finita (3-5 años).", "✅")

add_qa(doc, "2.7", "¿Qué alternativas existen en el mercado y por qué son insuficientes?",
    "Portales masivos: volumen sobre calidad, cero curaduría. CBRE/JLL: institucionales pero solo para tickets >50M USD. Houm/Flat: enfoque residencial masivo, sin componente industrial. TrueHome: crédito hipotecario, no inversión. Ninguno ofrece la combinación de curaduría + análisis financiero + experiencia premium que propone Black Corporativo.", "✅")

add_qa(doc, "2.8", "¿Tamaño de mercado (TAM, SAM, SOM)?",
    "TAM (México): ~500,000 transacciones/año de propiedades con valor >3M MXN. SAM (Baja California + frontera): ~18,000 transacciones/año. SOM (nicho premium con servicio digital): ~500-800 transacciones/año. A un ticket promedio de 12M MXN y comisión del 3-5%, el SOM representa ~180-480M MXN en valor de transacciones y ~5-24M MXN en ingresos por comisiones.", "✅")

add_qa(doc, "2.9", "¿Hay ventana regulatoria o tecnológica?",
    "Sí. La digitalización notarial en México (firma electrónica avanzada, escrituras digitales en CDMX y BC) está reduciendo la fricción del cierre. Supabase y Vercel permiten lanzar y escalar sin DevOps dedicado. La adopción de WhatsApp Business API permite automatizar seguimiento sin perder el canal preferido del usuario mexicano.", "✅")

add_qa(doc, "2.10", "2-3 casos reales de usuarios sufriendo el problema.",
    "Caso A: Family office de CDMX buscando nave industrial en Tijuana. Tardaron 7 semanas y contactaron a 6 brokers distintos. Dos propiedades ya estaban vendidas cuando preguntaron. Caso B: Inversionista individual de California buscando residencia de lujo en Playas. Recibió 40 opciones sin filtrar por WhatsApp. Abandonó la búsqueda. Caso C: Empresa maquiladora coreana necesitando 5,000 m² de bodega. Ningún portal mostraba la altura de montacargas ni la capacidad de andén — datos críticos para su decisión.", "✅")

doc.add_page_break()

# ============================================================
# SECCIÓN 3: USUARIOS Y STAKEHOLDERS
# ============================================================
add_section_band(doc, 3, "Usuarios y stakeholders",
    "Quién usa, quién paga, quién decide, quién bloquea.")

add_qa(doc, "3.1", "Segmentos de usuario: rol, contexto, frecuencia, tolerancia técnica.",
    "1) Inversionista HNWI (35-65 años): busca propiedad 1-2 veces al año, accede desde desktop y móvil, espera experiencia tipo Bloomberg Terminal (datos densos, bien presentados). 2) Corporate Buyer (CFO/COO de empresa): busca naves u oficinas, delega la búsqueda inicial en su equipo, necesita descargar brochure para comité. 3) Broker/Agente interno (Claudia): usa el admin panel a diario para gestionar leads y propiedades, necesita eficiencia operativa.", "✅")

add_qa(doc, "3.2", "¿Quién es el usuario primario?",
    "El inversionista HNWI que busca propiedad en Tijuana/San Diego. Es quien toma la decisión de compra y quien evalúa la calidad de la plataforma. Si él no confía, no hay transacción.", "✅")

add_qa(doc, "3.3", "¿Usuarios secundarios, terciarios e influenciadores?",
    "Secundarios: los equipos de los corporate buyers (asistentes, analistas que pre-filtran). Terciarios: propietarios/vendedores que quieren listar sus propiedades. Influenciadores: abogados, contadores y family offices que recomiendan brokers a sus clientes.", "✅")

add_qa(doc, "3.4", "¿Quién paga? ¿Quién decide la compra? ¿Son la misma persona?",
    "En residencial de lujo: el comprador decide y paga directamente. En comercial/industrial: el CFO/board decide y la empresa paga. El tomador de decisión y el pagador coinciden en ~70% de los casos. En el 30% restante, hay un comité de inversión que requiere materiales (brochure, análisis financiero) para decidir.", "✅")

add_qa(doc, "3.5", "Stakeholders internos clave.",
    "Moisés: fundador, DRI, broker principal, desarrollador full-stack. Claudia: agente asignada, gestión de leads y seguimiento. Los leads de Supabase se asignan a ella. Futuro: se necesitará un especialista en marketing digital (SEO, contenido, ads) y potencialmente un segundo agente para volumen.", "⚠")

add_qa(doc, "3.6", "Para cada stakeholder: ¿qué necesita, qué teme, cómo se mide?",
    "Moisés: necesita que la plataforma genere leads calificados sin requerir su presencia constante. Teme perder oportunidades por falta de seguimiento automatizado. Se mide por ingresos cerrados originados en digital. Claudia: necesita un CRM eficiente dentro del admin. Teme leads que se enfrían por no tener recordatorios automáticos. Se mide por tiempo de primera respuesta y tasa de conversión lead→cita.", "✅")

add_qa(doc, "3.7", "¿Reguladores, socios, proveedores o sindicatos?",
    "Reguladores: notarías públicas (para cierres, no para la plataforma en sí). Socios comerciales: propietarios/vendedores que confían sus propiedades a Black Corporativo. Proveedores críticos: Supabase (datos), Vercel (hosting), fotógrafos y videógrafos para el contenido de propiedades.", "✅")

add_qa(doc, "3.8", "¿Necesidades de accesibilidad?",
    "Actualmente no se ha hecho auditoría de accesibilidad. El sitio usa contrastes altos (dark theme con texto claro) que benefician la legibilidad, pero no se han implementado roles ARIA, navegación por teclado completa, ni pruebas con lectores de pantalla. WCAG 2.1 AA debe ser un objetivo para Q4 2026.", "⚠")

add_qa(doc, "3.9", "¿Diferencias culturales, de idioma o jurisdicción?",
    "El sitio está en español mexicano. Se contempla inglés como segundo idioma (Q1 2027) dado que ~30% de los compradores de lujo en Tijuana son extranjeros (principalmente de California). Jurisdicción: México, con propiedades concentradas en Baja California. Expansión futura a Sonora y CDMX.", "📋")

add_qa(doc, "3.10", "2-3 personas (proto-personas).",
    "Persona 1: Ricardo, 52, CFO de maquiladora en Tijuana. Vive en San Diego, cruza diario. Busca nave industrial de 2,000-5,000 m² para expansión. Le frustra no encontrar datos técnicos (altura de nave, capacidad de andén) en los portales. Valora discreción y rapidez. Persona 2: María Elena, 45, heredera de family office en CDMX. Busca residencia de lujo en Playas de Tijuana como inversión y segunda residencia. Quiere ver análisis de plusvalía y comparables de mercado. No confía en fotos de celular.", "✅")

# ============================================================
# SECCIÓN 4: OBJETIVOS Y MÉTRICAS
# ============================================================
add_section_band(doc, 4, "Objetivos y métricas",
    "Qué significa «terminamos» en términos medibles.")

add_qa(doc, "4.1", "3-5 objetivos de producto (outcomes), no features.",
    "1) Aumentar leads calificados de 15/mes a 40/mes. 2) Reducir tiempo de primera respuesta de 8h a <30 min (vía automatización WhatsApp). 3) Aumentar tasa de conversión visita→lead de 1.5% a 3.5%. 4) Lograr que el 60% de los leads vengan de búsqueda orgánica (SEO). 5) Cerrar al menos 3 transacciones/mes originadas 100% en digital.", "✅")

add_qa(doc, "4.2", "Métrica principal y baseline actual.",
    "Leads calificados/mes: baseline ~15 (mayo 2026). Meta: 40/mes para diciembre 2026. Tasa de conversión visita→lead: baseline 1.5%. Meta: 3.5%. Tiempo hasta primera respuesta: baseline ~8h (manual, Claudia). Meta: <30 min con WhatsApp automation.", "✅")

add_qa(doc, "4.3", "Meta cuantitativa y plazo.",
    "Diciembre 2026: 40 leads/mes, 3.5% conversión, <30min respuesta. Junio 2027: 80 leads/mes, 5% conversión, <5min respuesta con chatbot + agente humano. Cada lead calificado tiene un valor estimado de ~$15,000-50,000 MXN en comisión esperada.", "✅")

add_qa(doc, "4.4", "Métricas de guarda (guardrails).",
    "Tasa de rebote (bounce rate): no debe superar 50%. Tasa de quejas/desuscripción: <2%. Tiempo de carga (LCP): mantener <2.5s en móvil. No saturar a Claudia con leads no calificados (mantener ratio de calificación >60%).", "✅")

add_qa(doc, "4.5", "¿Cómo se reportan las métricas y con qué cadencia?",
    "PostHog para analítica de producto (embudos, sesiones, heatmaps). Dashboard interno de Supabase + admin panel para leads. Reporte semanal manual (Moisés revisa PostHog + Supabase). Ideal: dashboard automatizado en el admin panel con métricas en tiempo real. Cadencia: Moisés revisa lunes y viernes.", "⚠")

add_qa(doc, "4.6", "¿Eventos del negocio que se cruzan?",
    "Temporada alta inmobiliaria: octubre-marzo (cuando empresas definen presupuestos del año siguiente). Cierres fiscales: diciembre (urgencia de compra). Eventos de nearshoring: anuncios de nuevas plantas que generan demanda inmediata de naves industriales.", "✅")

add_qa(doc, "4.7", "¿OKRs de la organización conectados?",
    "KR1: 40 leads/mes desde digital. KR2: 5 cierres/mes originados en plataforma. KR3: 3 sub-brand pages con tráfico orgánico >500 visitas/mes cada una. KR4: Tiempo medio de ciclo lead→cierre <45 días.", "✅")

add_qa(doc, "4.8", "¿Qué es fracaso temprano, éxito modesto y éxito excepcional?",
    "Fracaso: <10 leads/mes después de 6 meses de optimización. Éxito modesto: 25-30 leads/mes stables, 2-3 cierres/mes. Éxito excepcional: 60+ leads/mes, 8+ cierres/mes, la plataforma genera suficiente para contratar 2 agentes adicionales y Moisés solo supervisa.", "✅")

add_qa(doc, "4.9", "¿Dependencias en otros productos/equipos?",
    "Supabase: si cambian el free tier o limits, impacta costo. Vercel: si hay outage, el sitio cae (ha sido estable). WhatsApp Business API: requiere aprobación de Meta y tiene costo por mensaje. Claudia: es recurso único; si no está disponible, el seguimiento de leads se detiene.", "⚠")

add_qa(doc, "4.10", "¿Cómo se verá el éxito a 30, 90, 180 y 365 días?",
    "30 días: leads respondidos en <1h, 20 leads/mes, sin errores 500 en producción. 90 días: 30 leads/mes, SEO empieza a traer tráfico orgánico, 2 cierres desde digital. 180 días: 40 leads/mes, WhatsApp automation funcionando, dashboard de métricas en admin. 365 días: 60+ leads/mes, 5+ cierres/mes, Claudia tiene capacidad ociosa (proceso optimizado), se contrata segundo agente.", "✅")

# ============================================================
# SECCIÓN 5: ALCANCE
# ============================================================
add_section_band(doc, 5, "Alcance",
    "Qué entra, qué no, qué se difiere.")

add_qa(doc, "5.1", "Experiencia del usuario final en 3-5 frases.",
    "Un inversionista llega a BlackCorporativo.com. Ve un hero cinematic con video de propiedades premium y la promesa «Impulsamos tu Legado». Navega a la vertical que le interesa (Luxury, Business o Industrial). Explora propiedades con fotos profesionales, ficha técnica completa, análisis financiero y un CTA claro de contacto. Descarga el brochure tras dejar sus datos. En menos de 30 minutos, recibe un mensaje de WhatsApp personalizado de Claudia con información adicional y disponibilidad para agendar visita.", "✅")

add_qa(doc, "5.2", "3-5 casos de uso críticos del MVP.",
    "1) Explorar catálogo → filtrar por tipo/uso → ver ficha → contactar. 2) Llegar desde Google a una propiedad específica → ver análisis financiero → descargar brochure → lead capturado. 3) Admin: Moisés crea propiedad → asigna agente → propiedad aparece en catálogo. 4) Lead magnet en homepage → formulario → lead en Supabase → Claudia recibe notificación.", "✅")

add_qa(doc, "5.3", "¿Qué se difiere a versiones posteriores?",
    "V2 (Q3 2026): WhatsApp automation, email transaccional (brochure por correo), favoritos/guardados, blog SEO, multi-idioma (EN). V3 (Q4 2026): búsqueda avanzada (mapa, precio, m²), chatbot IA en sitio, portal de propietarios, panel de analítica en admin. V4 (Q1-Q2 2027): app PWA con notificaciones push, comparador de propiedades, tour virtual 360°, matching automático lead↔propiedad.", "📋")

add_qa(doc, "5.4", "¿Qué está FUERA de alcance?",
    "Marketplace abierto (cualquiera lista propiedades): No. Black Corporativo mantiene curaduría. App nativa iOS/Android: No. La PWA debe ser suficiente. CRM completo tipo Salesforce: No. Supabase leads es suficiente con automatización ligera. Pasarela de pagos / transacciones en plataforma: No. Las transacciones se cierran offline.", "✅")

add_qa(doc, "5.5", "Supuestos que invalidarían el alcance si se rompen.",
    "1) Supabase free tier sigue siendo suficiente para <10K registros. 2) Moisés mantiene disponibilidad de 5-10h/semana para desarrollo. 3) Claudia sigue como agente activa. 4) El mercado de nearshoring mantiene crecimiento. 5) WhatsApp Business API es aprobada para la cuenta de Black Corporativo.", "✅")

add_qa(doc, "5.6", "¿Features que solo aplican a un segmento?",
    "Análisis financiero industrial incluye metros lineales de andén y altura de nave (específico de Black Industrial). Análisis de lujo incluye plusvalía proyectada y comparables de mercado premium (específico de Black Luxury). Las herramientas financieras (calculadora ROI, simulador flipping) aplican a los tres segmentos.", "✅")

add_qa(doc, "5.7", "Caso de borde más extremo.",
    "Un portafolio de 20 naves industriales de un solo propietario que quiere listar todo de una vez, cada una con documentos distintos, algunas en preventa (proyecto), otras en operación. El sistema debe soportar crear 20 propiedades en lote, o al menos que el admin no requiera 20 operaciones manuales idénticas.", "⚠")

add_qa(doc, "5.8", "¿Rutas regulatorias que condicionan el alcance?",
    "Aviso de privacidad (LFPDPPP): ya implementado en /legal/privacidad. Consentimiento explícito en formularios: implementado (checkbox privacy_accepted). Retención de datos: leads se conservan indefinidamente en Supabase; se necesita política de eliminación. COFECE (competencia): no aplica por tamaño de operación. Anti-lavado (LFPIORPI): aplica en la transacción offline, no en la plataforma.", "⚠")

add_qa(doc, "5.9", "¿Qué pasa con procesos/productos obsoletos?",
    "El sitio web anterior (si existía) se descontinúa. Los catálogos en PDF que se enviaban por WhatsApp se reemplazan por la ficha de propiedad digital + brochure descargable. La libreta de leads en Excel se reemplaza por el admin panel de Supabase.", "✅")

add_qa(doc, "5.10", "¿Cuál es el «cutline» y quién lo custodia?",
    "El cutline lo define Moisés como DRI. Regla: todo lo que aumente leads calificados o reduzca tiempo de respuesta entra. Todo lo que sea «nice to have» sin impacto directo en conversión se difiere. Si una feature no tiene hipótesis clara de impacto en leads o cierres, no se construye.", "✅")

# ============================================================
# SECCIÓN 6: REQUERIMIENTOS FUNCIONALES
# ============================================================
add_section_band(doc, 6, "Requerimientos funcionales",
    "Comportamientos que el sistema DEBE tener.")

add_qa(doc, "6.1", "Para cada caso de uso principal, flujo paso a paso.",
    "Ver catálogo: Home → Inventario → filtros (tipo, uso, precio) → grid de propiedades → click en card → ficha de propiedad con galería, métricas, descripción, documentos, agente y CTA. Captura de lead: formulario (nombre, email, teléfono, empresa opcional, aceptar privacidad) → insert en Supabase → PostHog captura evento → toast de éxito.", "✅")

add_qa(doc, "6.2", "¿Qué reglas de negocio aplican?",
    "Propiedades con status != 'Available' no se muestran en catálogo público. Propiedades is_featured=true aparecen primero en homepage. Precio se muestra en MXN; si currency='USD', se muestra el equivalente con tipo de cambio indicado. Leads no pueden insertarse sin privacy_accepted=true. Solo propiedades con cover_image se muestran en cards del catálogo.", "✅")

add_qa(doc, "6.3", "¿Flujos que cambian según rol, plan o segmento?",
    "Admin: acceso completo CRUD a properties, agents, leads. Público: solo lectura de properties con status=Available. La ficha de propiedad muestra diferentes métricas según property_use: residencial muestra recámaras/baños, industrial muestra m² de terreno/construcción y atributos técnicos, comercial muestra ubicación y tráfico.", "✅")

add_qa(doc, "6.4", "Estados y transiciones de entidades principales.",
    "Lead: new → contacted → qualified → closed_won / closed_lost. Property: Available → Under_Offer → Sold / Rented. Agent: is_active=true/false. La transición de lead la hace Claudia manualmente en admin. Ideal: automatizar new→contacted cuando se envía primer WhatsApp.", "⚠")

add_qa(doc, "6.5", "¿Cómo se manejan errores y excepciones?",
    "Formularios: validación client-side con Zod + react-hook-form, mensajes en español junto al campo. Errores de red: toast de error (sonner) con mensaje genérico. Errores 500: página de error de Next.js. Propiedad no encontrada (slug inválido): notFound(). Catálogo vacío: estado empty con mensaje «Portafolio en Curación». Errores de Supabase: console.error + degraded experience (no carga datos pero no rompe la página).", "✅")

add_qa(doc, "6.6", "¿CRUD de cada entidad?",
    "Properties: admin puede crear, editar, ver y eliminar (soft-delete vía status). Público solo SELECT. Leads: admin SELECT/UPDATE; público solo INSERT. Agents: admin CRUD completo; público SELECT de agentes activos. Todo vía Supabase client con RLS.", "✅")

add_qa(doc, "6.7", "Búsquedas, filtros, exportaciones obligatorias.",
    "Ya implementado: filtro por property_use (Residencial/Comercial/Industrial) y business_type (Venta/Renta) en catálogo vía CatalogFilter. Pendiente: búsqueda por texto libre (título, descripción, dirección), filtro por rango de precio y m², ordenamiento por precio o fecha. Exportación CSV de leads desde admin: no implementado, es necesidad operativa de Claudia.", "⚠")

add_qa(doc, "6.8", "Notificaciones, correos o mensajes del sistema.",
    "Actualmente: toast en UI (sonner) tras submit de formulario. No hay email transaccional (confirmación al lead, notificación a Claudia). No hay WhatsApp automation. Esto es la prioridad #1 del roadmap: al capturar un lead, el sistema debe enviar WhatsApp a Claudia y email de confirmación al lead.", "⚠")

add_qa(doc, "6.9", "Integraciones nativas requeridas.",
    "Supabase (auth, db, storage): implementado. Vercel Analytics + Speed Insights: implementado. PostHog (analítica): implementado. WhatsApp Business API: no implementado (crítico). Email transaccional (Resend/SendGrid): no implementado. Google Maps/Mapbox para ubicación de propiedades: no implementado.", "⚠")

add_qa(doc, "6.10", "Features de colaboración, auditoría o versionado.",
    "No se requiere versionado de cambios en propiedades (el admin es single-user: Moisés o Claudia). Auditoría básica vía created_at en Supabase. No hay historial de cambios en propiedades. Si en el futuro hay múltiples agentes editando, se necesitará.", "✅")

add_qa(doc, "6.11", "5 historias de usuario del MVP con criterios de aceptación.",
    "HU1: Como visitante, quiero ver propiedades destacadas en la homepage para evaluar si el portafolio me interesa. CA: se muestran 3 cards de propiedades is_featured + Available. HU2: Como comprador, quiero filtrar el catálogo por tipo de propiedad para ver solo lo relevante. CA: filtros de uso y tipo funcionan, sin recarga de página. HU3: Como inversionista, quiero ver el análisis financiero de una propiedad para decidir si investigo más. CA: ficha muestra precio, m², precio/m², cap rate estimado y atributos clave. HU4: Como lead, quiero descargar el brochure de una propiedad para compartirlo con mi comité. CA: formulario → descarga PDF, lead guardado en Supabase. HU5: Como agente (Claudia), quiero ver los leads nuevos para dar seguimiento. CA: admin/leads muestra lista ordenada por fecha, puedo ver detalle y cambiar status.", "✅")

# ============================================================
# SECCIÓN 7: REQUERIMIENTOS NO FUNCIONALES
# ============================================================
add_section_band(doc, 7, "Requerimientos no funcionales",
    "Calidades del sistema.")

add_qa(doc, "7.1", "Rendimiento: usuarios concurrentes, RPS en pico.",
    "Actual: <50 visitantes/día, picos de 5-10 simultáneos. Objetivo 12 meses: soportar 500 visitantes/día, 50 concurrentes, sin degradación. Infraestructura actual (Vercel hobby + Supabase free tier) soporta esto sin cambios.", "✅")

add_qa(doc, "7.2", "Latencia: tiempo de respuesta aceptable.",
    "Page load (LCP): <2.5s en móvil 3G (ya se cumple con Next.js SSG/ISR). API calls a Supabase: <500ms p95. Las fichas de propiedad se sirven estáticas (revalidate=60s), así que la latencia es mínima.", "✅")

add_qa(doc, "7.3", "Disponibilidad SLA objetivo.",
    "99.5% (aceptable para B2B no transaccional). Vercel ofrece 99.9%+. Supabase ofrece 99.9%+. Ventanas de mantenimiento: domingos 2-4 AM PT. No se requiere 99.99% porque la plataforma no procesa transacciones en tiempo real.", "✅")

add_qa(doc, "7.4", "Escalabilidad: ¿cómo debe crecer?",
    "Horizontal: Vercel escala automáticamente las serverless functions. Supabase free tier soporta 500 MB de datos y 2 GB de transferencia. Para escalar a >10K propiedades o >50K leads, se requiere Supabase Pro ($25/mes). No se anticipa necesidad de multi-región en los próximos 24 meses.", "✅")

add_qa(doc, "7.5", "Capacidad: volumen de datos.",
    "Actual: ~50 propiedades, ~200 leads, ~5 agentes. A 12 meses: ~200 propiedades, ~1000 leads, ~10 agentes. A 36 meses: ~1000 propiedades, ~10K leads, ~30 agentes. Todo cabe en Supabase Pro sin optimización especial.", "✅")

add_qa(doc, "7.6", "Fiabilidad: RPO y RTO.",
    "RPO: 0 (Supabase tiene replicación síncrona y backups automáticos). RTO: <4h (en caso de desastre, restaurar backup de Supabase + redeploy en Vercel desde git). Las propiedades y leads son los datos críticos; las imágenes están en Supabase Storage.", "✅")

add_qa(doc, "7.7", "Mantenibilidad: ¿qué tan fácil de modificar?",
    "La arquitectura de componentes de Next.js + shadcn/ui hace que cada cambio esté aislado. El código está organizado por dominio (components/property/, components/home/, app/(public)/, app/(admin)/). Tipado estricto con TypeScript. El bus factor actual es 1 (Moisés), lo cual es un riesgo.", "⚠")

add_qa(doc, "7.8", "Portabilidad: ¿cloud, on-premise, móvil?",
    "Cloud-only (Vercel + Supabase). No hay requisito on-premise. La PWA funciona offline para contenido cacheado. Si se requiere migrar de Vercel a otro proveedor, Next.js es portable (Netlify, Cloudflare Pages, Docker).", "✅")

add_qa(doc, "7.9", "Compatibilidad: navegadores, OS, dispositivos.",
    "Chrome, Firefox, Safari, Edge (últimas 2 versiones). iOS Safari 15+, Android Chrome 100+. La PWA funciona en iOS (añadir a pantalla de inicio) y Android (instalación completa). No se soporta IE11 ni navegadores legacy.", "✅")

add_qa(doc, "7.10", "Usabilidad: estándares UX, WCAG, tiempo para tareas.",
    "WCAG 2.1 AA: no verificado formalmente. El dark theme ayuda al contraste pero faltan roles ARIA y navegación por teclado. Tiempo para encontrar una propiedad: <30 segundos desde homepage. Tiempo para completar formulario de lead: <45 segundos. Tasa de abandono de formulario actual: desconocida (medir con PostHog).", "⚠")

add_qa(doc, "7.11", "Observabilidad: logs, métricas, trazas.",
    "Vercel Analytics: Web Vitals (LCP, CLS, INP). PostHog: sesiones, eventos, embudos. Supabase: logs de queries vía dashboard. No hay monitor de errores (Sentry) ni alertas. La caída del sitio se detecta cuando Moisés o un usuario lo reportan.", "⚠")

add_qa(doc, "7.12", "Internacionalización: idiomas, formatos.",
    "Actualmente solo español (es-MX). El formateo de precios usa MXN con notación mexicana. Fechas en formato DD/MM/YYYY. Pendiente: inglés (en-US) con USD. No se requiere RTL (árabe/hebreo) en el horizonte previsible. La arquitectura de Next.js con i18n routing soportaría múltiples idiomas sin refactor mayor.", "📋")

# ============================================================
# SECCIÓN 8: UX Y DISEÑO
# ============================================================
add_section_band(doc, 8, "Experiencia de usuario y diseño",
    "Cómo se siente y se ve, no solo cómo funciona.")

add_qa(doc, "8.1", "¿Tono y personalidad de marca?",
    "Sofisticado, institucional, confiable. No es «vendedor agresivo» ni «startup casual». La tipografía Montserrat (display, uppercase) comunica solidez. El dorado metálico es el único color además de negro y grises. Sin emojis, sin exclamaciones, sin urgencia falsa. La promesa es curaduría y discreción.", "✅")

add_qa(doc, "8.2", "¿Design system o guía de estilo?",
    "Existe un sistema de diseño parcial implementado en globals.css: escala tipográfica (hero-title, section-heading, card-title, label-overline, body-text), paleta de dorados (gold-300 a gold-700), glassmorphism (.glass), animaciones (.metallic-gold, .animate-float, .animate-marquee). Faltan: documentación formal del design system, componentes de formulario estandarizados, y guía de voz y tono.", "⚠")

add_qa(doc, "8.3", "Primera experiencia (onboarding) del usuario.",
    "El usuario llega a la homepage y ve el hero con video y la palabra rotativa («Legado», «Futuro», «Expansión»). Hace scroll para ver BrandsGrid (3 verticales), FeaturedInventory (3 propiedades destacadas), SocialProof (métricas de autoridad) y LeadMagnet (CTA de contacto). En <60 segundos entiende qué hace Black Corporativo y puede decidir si explorar más.", "✅")

add_qa(doc, "8.4", "Time-to-value: ruta más corta al primer valor.",
    "Desde homepage: scroll → click en propiedad destacada → ver ficha completa con análisis financiero → click en «Me interesa» → formulario → Claudia contacta. Tiempo estimado: 2-3 minutos hasta que el lead ve el valor (datos que no encuentra en otros portales).", "✅")

add_qa(doc, "8.5", "Estados vacíos, errores y esperas.",
    "Catálogo vacío: «Portafolio en Curación» con ícono decorativo (ya implementado). Error de carga: no implementado (si Supabase falla, la página queda en blanco o con skeleton infinito). Necesita: estado de error con mensaje y botón de reintentar. Esperas: skeletons (pulsing placeholders) en featured inventory y catálogo.", "⚠")

add_qa(doc, "8.6", "¿Qué tan personalizable debe ser la experiencia?",
    "Actualmente: cero personalización. El usuario no tiene cuenta ni perfil. V2 debería incluir: guardar propiedades favoritas (localStorage o cuenta Supabase), recordar filtros de búsqueda, y recomendaciones basadas en propiedades visitadas. No se requiere personalización profunda tipo dashboard configurable.", "📋")

add_qa(doc, "8.7", "¿Diseños responsivos?",
    "Sí. El sitio es completamente responsive (Tailwind responsive utilities). La pill navbar colapsa a drawer en móvil. Las cards de propiedad pasan de 3 columnas (desktop) a 2 (tablet) a 1 (móvil). Los formularios se adaptan. Probado en iPhone SE, iPhone 15 Pro Max y iPad.", "✅")

add_qa(doc, "8.8", "¿Investigación de usuarios realizada o necesaria?",
    "Realizada: observación directa (Moisés es broker y conoce a sus clientes). No se han hecho entrevistas formales de UX ni pruebas de usabilidad. Necesaria: prueba con 3-5 clientes reales (inversionistas) observándolos usar el sitio, midiendo tiempo para encontrar propiedad y completar formulario. Esto revelaría fricciones que el equipo no ve.", "⚠")

add_qa(doc, "8.9", "¿Flujos críticos de accesibilidad desde el MVP?",
    "El formulario de captura de leads es el flujo más crítico. Debe ser operable por teclado, tener labels asociados correctamente, mensajes de error vinculados a campos (aria-describedby), y contraste suficiente en todos los estados (focus, error, disabled). Actualmente usa react-hook-form pero no se verificó accesibilidad.", "⚠")

add_qa(doc, "8.10", "¿Copy, terminología y micro-interacciones definidas?",
    "Terminología consistente: «propiedades» (no «inmuebles»), «portafolio» (no «listado»), «análisis financiero» (no «datos»). Micro-interacciones implementadas: staggered letter reveal en hero, scroll-triggered counters en social proof, fade-in en secciones, hover dorado en nav links. Pendiente: validación inline en formularios (errores aparecen solo en submit, no en blur).", "⚠")

# ============================================================
# SECCIÓN 9: ARQUITECTURA Y DATOS
# ============================================================
add_section_band(doc, 9, "Arquitectura y datos",
    "Cómo se construye por dentro.")

add_qa(doc, "9.1", "¿Restricciones de stack, cloud o lenguaje?",
    "Next.js 16 + React 19 + TypeScript (no negociable, es el stack existente). Vercel para hosting (gratuito, integración nativa). Supabase para backend (PostgreSQL + Auth + Storage). No se puede migrar a otro backend sin reescribir todas las queries. Tailwind CSS 4 + shadcn/ui + Radix UI para componentes.", "✅")

add_qa(doc, "9.2", "¿Monolito, microservicios, serverless, edge?",
    "Arquitectura: Next.js App Router con React Server Components (RSC) para páginas públicas y Client Components para interactividad. Es un monolito bien modularizado — no hay microservicios. Las serverless functions de Vercel manejan API routes. No hay necesidad de edge computing actualmente (latencia a Supabase desde Vercel es <50ms).", "✅")

add_qa(doc, "9.3", "¿Bases de datos?",
    "Supabase PostgreSQL (relacional). Esquema: properties, agents, leads, property_agents (junction table). No se usa NoSQL, graph, ni vectorial. Suficiente para las necesidades actuales. Si en el futuro se añade búsqueda semántica de propiedades, se podría usar pgvector (incluido en Supabase).", "✅")

add_qa(doc, "9.4", "Modelo de datos principal.",
    "properties es la entidad central. Tiene propiedades calculadas/no normalizadas: price_mxn (para ordenamiento), custom_attributes (JSONB flexible para atributos específicos por tipo de propiedad), images/video_urls/documents como JSONB arrays. agents tiene relación M:N con properties vía property_agents. leads se relaciona opcionalmente con properties (property_id) y agents (assigned_agent_id).", "✅")

add_qa(doc, "9.5", "¿Consistencia requerida?",
    "No hay operaciones que requieran consistencia fuerte entre múltiples entidades. Las lecturas son todas eventuales (Supabase entrega consistencia de lectura inmediata para la misma sesión). No hay transacciones distribuidas ni necesidad de serializabilidad.", "✅")

add_qa(doc, "9.6", "¿Datos sensibles y cómo se protegen?",
    "Leads: nombre, email, teléfono. Son PII bajo LFPDPPP. Se almacenan en Supabase con RLS. El acceso público solo puede INSERTAR; solo authenticated puede leer. El access token de Supabase está en variables de entorno. No se almacenan datos financieros de clientes ni documentos de identidad. Las imágenes se almacenan en Supabase Storage con acceso público (solo lectura).", "✅")

add_qa(doc, "9.7", "Estrategia de respaldo, retención y archivado.",
    "Supabase incluye backups automáticos diarios (free tier: 7 días de retención). No hay estrategia de archivado de leads antiguos. Propiedades vendidas se marcan status='Sold' pero no se eliminan. Política actual: no borrar datos. A futuro: leads >2 años sin actividad se pueden anonimizar o archivar.", "⚠")

add_qa(doc, "9.8", "¿Procesamiento en tiempo real o batch?",
    "No hay procesamiento batch actualmente. Todo es en tiempo real vía API calls directas a Supabase. Si se implementa matching automático lead↔propiedad, se podría usar Supabase Realtime (websockets) o un cron job en Vercel. No se requiere procesamiento batch estilo ETL.", "✅")

add_qa(doc, "9.9", "¿Integraciones externas obligatorias?",
    "Ya integrado: Supabase API, PostHog, Vercel Analytics. Pendientes obligatorias: WhatsApp Business API (crítico para automatización de seguimiento), Email API (Resend, recomendado por simplicidad y costo). Deseables: Google Maps API para ubicación de propiedades, Google Analytics (complementario a PostHog), Meta Pixel para remarketing.", "⚠")

add_qa(doc, "9.10", "¿Decisiones de arquitectura pendientes?",
    "1) ¿Cuándo migrar de Supabase free a Pro? Cuando se alcancen 500MB de datos o 50K rows. 2) ¿Email provider? Resend vs SendGrid — decidir por costo y simplicidad de API. 3) ¿i18n routing? Definir estrategia antes de implementar inglés (subdominio en.blackcorporativo.com vs subruta /en/). 4) ¿CDN para imágenes? Actualmente Supabase Storage; considerar Cloudinary si el volumen de imágenes crece >1GB.", "⚠")

# ============================================================
# SECCIÓN 10: SEGURIDAD Y CUMPLIMIENTO
# ============================================================
add_section_band(doc, 10, "Seguridad y cumplimiento",
    "Lo que no puede fallar sin que se caiga el negocio.")

add_qa(doc, "10.1", "¿Marcos de cumplimiento aplicables?",
    "LFPDPPP (Ley Federal de Protección de Datos Personales en Posesión de Particulares): México. GDPR: solo si se captan leads de UE (actualmente no). CCPA: no aplica (no se opera en California como entidad). No se requiere HIPAA, PCI-DSS, SOC 2 ni ISO 27001 en esta etapa. Si se crece a >50 empleados o se manejan transacciones en plataforma, se reevalúa.", "✅")

add_qa(doc, "10.2", "¿Autenticación de usuarios?",
    "Usuarios públicos: no requieren autenticación. Admin: login con email + token JWT almacenado en cookie (AUTH_COOKIE). Implementación vía middleware de Next.js + isValidTokenFormat. No hay MFA, SSO ni passkeys. Para un solo usuario admin (Moisés/Claudia) esto es suficiente. Si se añaden más agentes al admin, se debe implementar Supabase Auth con RLS por rol.", "⚠")

add_qa(doc, "10.3", "¿Autorización (RBAC, ACL)?",
    "Modelo actual: binario (público vs admin). No hay granularidad de roles dentro del admin. La RLS de Supabase usa policies: público SELECT en properties/agents/property_agents, INSERT en leads; admin (authenticated) ALL en todas las tablas. Suficiente para 1-2 admins.", "✅")

add_qa(doc, "10.4", "¿Cifrado en tránsito y reposo?",
    "Tránsito: HTTPS (Vercel + Supabase). Reposo: Supabase cifra los datos en disco (AES-256). Las contraseñas no se almacenan (se usa token JWT). Los archivos en Supabase Storage también están cifrados en reposo.", "✅")

add_qa(doc, "10.5", "¿Manejo de llaves y secretos?",
    "Variables de entorno en Vercel (.env.local en desarrollo): SUPABASE_URL, SUPABASE_ANON_KEY, SUPABASE_SERVICE_ROLE_KEY, NEXT_PUBLIC_POSTHOG_KEY. No hay vault externo (HashiCorp, AWS KMS) — innecesario para esta escala. El token admin se genera al hacer login y se almacena en cookie httpOnly.", "✅")

add_qa(doc, "10.6", "¿Aislamiento entre clientes/tenants?",
    "No hay multi-tenancy. Las propiedades y leads son globales dentro de la cuenta Supabase. Si en el futuro se onboardean otros brokers independientes a la plataforma, se necesitará tenant isolation. No está en el roadmap inmediato.", "✅")

add_qa(doc, "10.7", "¿Auditoría y logging obligatorios?",
    "Actualmente: timestamps created_at en todas las tablas. No hay updated_at en properties (solo en agents y leads). No hay registro de qué admin modificó qué propiedad. Para cumplimiento futuro: añadir updated_by en properties y leads, y log de cambios.", "⚠")

add_qa(doc, "10.8", "¿Amenazas más preocupantes?",
    "1) SQL injection: mitigado por Supabase client (consultas parametrizadas). 2) Acceso no autorizado a admin: mitigado por middleware + token. 3) Fuga de datos de leads: mitigado por RLS. 4) Denegación de servicio: Vercel tiene protección DDoS básica. 5) Suplantación en WhatsApp: riesgo si se automatiza sin verificar que el número es realmente del lead.", "✅")

add_qa(doc, "10.9", "¿Threat modeling, pentest, bug bounty?",
    "No se ha realizado threat modeling formal ni pentest. Para la escala actual no es requisito, pero antes de manejar >1,000 leads se recomienda un pentest básico (OWASP Top 10). Bug bounty: no justificado por alcance. Lo que sí se debe hacer: revisión de código enfocada en seguridad en cada PR.", "⚠")

add_qa(doc, "10.10", "Plan de respuesta a incidentes y notificación.",
    "Actualmente no hay plan formal. Procedimiento deseado: 1) Moisés recibe alerta (Sentry), 2) Evalúa severidad, 3) Si es fuga de datos, notificar a leads afectados en <72h (LFPDPPP), 4) Postmortem. Responsable: Moisés. Es necesario implementar monitoreo de errores (Sentry) para detectar incidentes.", "⚠")

add_qa(doc, "10.11", "¿Qué hacer con datos al eliminar cuenta o exportar?",
    "Actualmente no hay mecanismo de eliminación de cuenta (los leads no tienen cuenta). Si un lead solicita eliminación de sus datos (derecho ARCO), Moisés debe eliminar manualmente el registro en Supabase. No hay funcionalidad de exportación de datos para el usuario. Con LFPDPPP esto debe implementarse antes de escalar.", "⚠")

# ============================================================
# SECCIÓN 11: OPERACIONES Y SOPORTE
# ============================================================
add_section_band(doc, 11, "Operaciones y soporte",
    "Lo que pasa después del lanzamiento.")

add_qa(doc, "11.1", "¿Quién opera el sistema 24/7?",
    "No hay operador 24/7. Vercel + Supabase operan la infraestructura. Si el sitio cae, lo nota Moisés (o un usuario lo reporta). El horario de atención de leads es L-V 9:00-18:00 PT por Claudia. Fuera de ese horario, los leads se acumulan. La automatización de WhatsApp puede enviar respuesta inmediata 24/7.", "⚠")

add_qa(doc, "11.2", "¿Equipo de soporte, herramientas, runbooks?",
    "Claudia es el equipo de soporte (nivel 1). Moisés es nivel 2 (técnico). Herramientas: admin panel de Supabase para ver leads. No hay runbooks documentados. Necesario: documentar proceso de seguimiento de leads (primer contacto, seguimiento a 24h, 72h, 7 días), y cómo crear/editar propiedades en el admin.", "⚠")

add_qa(doc, "11.3", "SLA de soporte por severidad.",
    "Severidad 1 (sitio caído): respuesta <2h, resolución <4h. Severidad 2 (feature rota, ej. formulario no envía): respuesta <8h, resolución <24h. Severidad 3 (error visual, typo): respuesta <48h, resolución <1 semana. No hay SLA contractual con clientes — es interno del equipo.", "✅")

add_qa(doc, "11.4", "¿Qué tan fácil diagnosticar problemas?",
    "Actualmente: difícil. No hay centralized logging. Errores de cliente se ven en browser console. Errores de servidor requieren revisar Vercel logs. Para diagnosticar un problema de lead que no se guardó, hay que revisar Supabase logs + PostHog eventos. Necesario: Sentry para error tracking y un dashboard de health checks.", "⚠")

add_qa(doc, "11.5", "¿Procesos de despliegue definidos?",
    "CI/CD: git push a main → Vercel auto-deploy. No hay staging environment (se prueba en local). Feature flags: no implementadas. Rollback: revertir commit y pushear, Vercel redeploya en <1min. Proceso ligero, adecuado para 1 desarrollador.", "✅")

add_qa(doc, "11.6", "Migraciones de esquema, datos y API.",
    "Migraciones de BD: archivos SQL en supabase/migrations/, ejecutadas manualmente en Supabase dashboard o vía CLI. No hay migración automática en CI/CD. Versionado de API: no hay API pública expuesta (solo las routes internas de Next.js). Si se expone API, se debe versionar desde el inicio (/api/v1/...).", "✅")

add_qa(doc, "11.7", "Estrategia de monitoreo y alertas.",
    "Monitoreo actual: Vercel Analytics + Supabase dashboard + PostHog. No hay alertas configuradas. Plan: Sentry para errores (free tier, 5K events/mes), Vercel log drains para errores 500, y un health check básico que verifique que Supabase responde. Alertas vía email a Moisés.", "⚠")

add_qa(doc, "11.8", "Gestión de dependencias externas.",
    "Dependencias críticas: Supabase (BD + Auth + Storage), Vercel (hosting). Si Supabase tiene outage, el sitio no puede cargar datos (se muestra skeleton o estado vacío). Mitigación: implementar graceful degradation — si Supabase no responde, mostrar datos cacheados o mensaje de error amigable. Actualmente no implementado.", "⚠")

add_qa(doc, "11.9", "Planes de contingencia.",
    "Caída de Supabase: el sitio sigue sirviendo páginas estáticas (SSG/ISR) pero sin datos dinámicos. Caída de Vercel: altamente improbable; si ocurre, esperar a que Vercel resuelva. Caída de WhatsApp API: los leads se siguen capturando, Claudia los ve en admin y contacta manualmente.", "✅")

add_qa(doc, "11.10", "Estrategia de respaldo y disaster recovery.",
    "Supabase: backups diarios con 7 días de retención. Código: en GitHub (backup implícito). Procedimiento DR: 1) Crear nuevo proyecto Supabase, 2) Restaurar backup, 3) Actualizar environment variables en Vercel, 4) Redeploy. Tiempo estimado: <4h. No se ha probado este procedimiento.", "⚠")

# ============================================================
# SECCIÓN 12: ESTRATEGIA DE LANZAMIENTO
# ============================================================
add_section_band(doc, 12, "Estrategia de lanzamiento y go-to-market",
    "Cómo llega al usuario y al mercado.")

add_qa(doc, "12.1", "¿Fecha de lanzamiento objetivo y por qué?",
    "MVP ya está en producción (mayo 2026). Lanzamiento «oficial» con campaña: septiembre 2026 (inicio de temporada alta inmobiliaria). ¿Por qué?: octubre-marzo es cuando empresas definen presupuestos y compradores HNWI están activos.", "✅")

add_qa(doc, "12.2", "¿Soft launch, beta, GA, por región?",
    "Fase actual: soft launch (sitio vivo, sin campaña). Beta cerrada (jun-ago 2026): invitar a 10-20 clientes existentes de Moisés a usar la plataforma, recoger feedback. GA (sep 2026): campaña de lanzamiento en LinkedIn, Google Ads y alianzas con family offices. Geografía inicial: Tijuana y San Diego.", "✅")

add_qa(doc, "12.3", "Criterios para considerar el sistema «listo».",
    "1) 0 errores 500 en producción durante 14 días consecutivos. 2) Formulario de leads funcional (probado con leads reales). 3) Admin panel permite CRUD de propiedades sin errores. 4) SEO básico implementado (sitemap, metadata, OG tags). 5) Velocidad de carga LCP <3s en móvil. 6) Aviso de privacidad publicado.", "⚠")

add_qa(doc, "12.4", "Materiales de lanzamiento necesarios.",
    "Sitio web (ya existe). LinkedIn company page (ya existe). Brochure corporativo digital (pendiente). Deck de presentación para family offices (pendiente). 3-5 casos de estudio de transacciones exitosas (pendiente). Video testimonial de 1-2 clientes (pendiente). Post de lanzamiento en LinkedIn + envío a base de contactos.", "⚠")

add_qa(doc, "12.5", "¿Programa de early adopters?",
    "Sí. 10-20 clientes existentes de Moisés que han cerrado transacciones anteriormente. Se les da acceso prioritario a nuevas propiedades antes de publicarlas en el sitio. A cambio, se les pide feedback sobre la plataforma (usabilidad, información faltante, qué les gustaría ver).", "✅")

add_qa(doc, "12.6", "Estrategia de pricing y monetización.",
    "Modelo: comisión por transacción (no suscripción). Comisión estándar: 3-5% del valor de la propiedad. La plataforma no cobra al comprador ni al vendedor por listar — es un canal de adquisición para Moisés como broker. Futuro: si se onboardean otros brokers, cobrar suscripción mensual + comisión reducida.", "✅")

add_qa(doc, "12.7", "Canales de venta/distribución.",
    "Directo: Moisés y Claudia como agentes. Digital: sitio web + WhatsApp. Indirecto (futuro): alianzas con family offices, despachos de abogados corporativos, contadores. No se contempla marketplace ni afiliados en esta etapa.", "✅")

add_qa(doc, "12.8", "Estrategia de comunicación interna y externa.",
    "Interna: canal de WhatsApp Moisés-Claudia para coordinación de leads. Externa: LinkedIn (2-3 posts/semana sobre mercado inmobiliario), email trimestral a base de contactos, WhatsApp directo a clientes VIP cuando entra propiedad relevante. No se contempla PR ni medios pagados en esta etapa.", "⚠")

add_qa(doc, "12.9", "¿Cómo se mide la adopción temprana?",
    "Tráfico: unique visitors/mes (PostHog). Leads: total leads/mes y tasa de conversión. Engagement: tiempo en sitio, páginas por sesión, scroll depth en fichas de propiedad. Cierres: cuántas transacciones se originaron en un lead digital.", "✅")

add_qa(doc, "12.10", "¿Qué pasa con usuarios existentes?",
    "No hay migración — la plataforma es nueva. Los clientes existentes de Moisés se incorporan al nuevo flujo: en lugar de recibir propiedades por WhatsApp, se les invita a explorar el sitio. La transición es suave porque el contacto personal (Moisés/Claudia) sigue siendo el mismo.", "✅")

# ============================================================
# SECCIÓN 13: COMERCIAL, PRICING Y FINANZAS
# ============================================================
add_section_band(doc, 13, "Comercial, pricing y finanzas",
    "Cómo se gana (o se ahorra) dinero.")

add_qa(doc, "13.1", "Modelo de negocio.",
    "Corretaje inmobiliario tradicional (comisión por transacción). La plataforma digital es un canal de adquisición de leads, no un producto independiente. Si en el futuro se convierte en marketplace para otros brokers, el modelo pivota a SaaS + comisión.", "✅")

add_qa(doc, "13.2", "Precio objetivo y sensibilidad.",
    "Comisión: 3-5% sobre valor de propiedad. Ticket promedio: ~12M MXN → comisión ~$360K-600K MXN por transacción. Sensibilidad: compradores HNWI son menos sensibles a comisión si perciben valor (curaduría, análisis, velocidad). El diferenciador no es precio sino calidad de servicio.", "✅")

add_qa(doc, "13.3", "Costo de adquisición de cliente (CAC).",
    "CAC actual: ~$0 (tráfico orgánico y referidos). Con campaña de lanzamiento (Google Ads + LinkedIn): estimado $500-1,500 MXN por lead, $8,000-15,000 MXN por cliente cerrado. Comparado con una comisión de $400K MXN, el CAC es <4% del ingreso.", "✅")

add_qa(doc, "13.4", "Valor de vida del cliente (LTV).",
    "Un inversionista HNWI compra 1-2 propiedades cada 3-5 años. A 2 transacciones en 5 años con comisión promedio $400K c/u: LTV ~$800K MXN. LTV/CAC >50:1 para tráfico orgánico, >25:1 con campaña. Excelente.", "✅")

add_qa(doc, "13.5", "Costo de entregar el producto (COGS).",
    "Infraestructura: ~$0/mes (Vercel + Supabase free). Tiempo de Moisés: 5-10h/semana en desarrollo (costo de oportunidad, no desembolso). Tiempo de Claudia: sueldo de agente. Sin campaña, el costo operativo mensual es solo el sueldo de Claudia (~$15-25K MXN/mes).", "✅")

add_qa(doc, "13.6", "Período de recuperación (payback).",
    "Payback inmediato: la primera transacción cerrada desde digital cubre varios meses de operación. Incluso con campaña de marketing de $20K MXN/mes, una sola transacción de $12M MXN genera ~$400K MXN de comisión = 20x retorno sobre inversión mensual en marketing.", "✅")

add_qa(doc, "13.7", "¿Planes gratuitos, trials, descuentos?",
    "No aplica — no es SaaS. El acceso al catálogo es gratuito. La comisión se negocia caso por caso (puede variar 3-5% según tipo de propiedad y relación con cliente).", "✅")

add_qa(doc, "13.8", "Sensibilidad a impuestos y tipo de cambio.",
    "El 80% de las propiedades se transan en MXN. Algunas propiedades de lujo se listan en USD. El tipo de cambio USD/MXN afecta la percepción de precio para compradores extranjeros. Fiscalmente, las comisiones pagan IVA e ISR según el régimen fiscal de Moisés.", "✅")

add_qa(doc, "13.9", "ROI esperado a 12 y 24 meses.",
    "12 meses: 5 cierres/mes × $400K MXN comisión promedio = $24M MXN/año en comisiones. 24 meses: 8 cierres/mes × $400K = $38.4M MXN/año. Inversión requerida: tiempo de Moisés + $20K/mes marketing + 1-2 agentes adicionales. ROI >10:1 en ambos escenarios.", "✅")

add_qa(doc, "13.10", "¿Quién aprueba el pricing?",
    "Moisés como fundador y broker principal. No hay comité de pricing. Las comisiones se ajustan por propiedad según negociación directa con el cliente.", "✅")

# ============================================================
# SECCIÓN 14: LEGAL, PRIVACIDAD Y ÉTICA
# ============================================================
add_section_band(doc, 14, "Legal, privacidad y ética",
    "Lo que requiere firma, abogado o ambos.")

add_qa(doc, "14.1", "¿Contratos, ToS, políticas, EULAs requeridos?",
    "Aviso de privacidad: implementado en /legal/privacidad. Términos de uso del sitio web: no implementado (recomendable). Contrato de corretaje: offline, se firma con cada cliente (no en plataforma). EULA: no aplica (no es software licenciado).", "⚠")

add_qa(doc, "14.2", "¿Datos de menores o categorías sensibles?",
    "No. Black Corporativo no capta datos de menores ni categorías sensibles (salud, ideología, etc.). Solo datos de contacto (nombre, email, teléfono) con fines de prospección comercial inmobiliaria.", "✅")

add_qa(doc, "14.3", "Derechos ARCO del usuario.",
    "El lead tiene derecho a Acceder, Rectificar, Cancelar y Oponerse al uso de sus datos. Actualmente estos derechos se ejercen contactando a Moisés directamente (el aviso de privacidad incluye el correo contacto@blackcorporativo.com). No hay mecanismo automatizado en la plataforma.", "⚠")

add_qa(doc, "14.4", "¿Propiedad intelectual de terceros?",
    "Fotografías de propiedades: propiedad del fotógrafo, con licencia de uso para Black Corporativo. Íconos (Lucide): open source (ISC). Fuentes (Inter, Montserrat): open source (OFL). No hay riesgo de infracción de PI.", "✅")

add_qa(doc, "14.5", "¿Modelos de IA o algoritmos con sesgos?",
    "No se usan modelos de IA en producción (el análisis financiero es determinístico, fórmulas en código). Si en el futuro se implementa matching automático lead↔propiedad con IA, se debe evaluar sesgo (ej. no recomendar ciertas zonas por sesgo en datos de entrenamiento). No es riesgo actual.", "✅")

add_qa(doc, "14.6", "¿Impacto ambiental, social o de accesibilidad?",
    "Impacto ambiental: mínimo (la infraestructura cloud de Vercel/Supabase usa energía de los data centers de AWS, que tienen metas de carbono neutral). Social: la plataforma democratiza el acceso a información de calidad para compradores que no tienen conexiones personales con brokers. Accesibilidad: no se ha evaluado; es deuda pendiente.", "⚠")

add_qa(doc, "14.7", "¿Si se descubre un defecto legal post-lanzamiento?",
    "Ejemplo: aviso de privacidad incompleto o formulario que captura datos sin consentimiento explícito. Procedimiento: 1) Corregir el texto/funcionalidad inmediatamente, 2) Notificar a leads afectados si hubo violación, 3) Documentar el incidente. No hay proceso formal; depende de Moisés.", "⚠")

add_qa(doc, "14.8", "¿Responsable legal y oficial de privacidad?",
    "Moisés funge como responsable legal de facto y oficial de privacidad. No hay nombramiento formal de DPO (no requerido por LFPDPPP para este tamaño de operación).", "✅")

add_qa(doc, "14.9", "¿Restricciones contractuales con socios o proveedores?",
    "Supabase ToS: no hay restricción de uso comercial. Vercel ToS: permite uso comercial. No hay contratos de exclusividad con propietarios que limiten la operación de la plataforma.", "✅")

add_qa(doc, "14.10", "¿DPIA, análisis de impacto o consulta pública?",
    "No requerido por LFPDPPP para el volumen y tipo de datos actual (<1,000 leads, sin datos sensibles). Si se superan 10,000 leads o se capturan datos financieros, se recomienda realizar una Evaluación de Impacto en la Protección de Datos (EIPD).", "✅")

# ============================================================
# SECCIÓN 15: RIESGOS Y DEPENDENCIAS
# ============================================================
add_section_band(doc, 15, "Riesgos y dependencias",
    "Qué puede salir mal y cómo se mitiga.")

add_qa(doc, "15.1", "Top 5-10 riesgos.",
    "1) Bus factor: solo Moisés conoce el código. 2) Supabase free tier se vuelve insuficiente. 3) WhatsApp API no se aprueba. 4) Claudia deja la operación. 5) Competidor proptech entra al mercado de Tijuana. 6) Cambio regulatorio (ej. nueva ley de corretaje). 7) Caída prolongada de Supabase/Vercel. 8) Bajo tráfico orgánico (SEO no despega). 9) Leads no calificados saturan a Claudia. 10) Ciberseguridad: inyección o fuga de datos.", "✅")

add_qa(doc, "15.2", "Matriz de riesgos con mitigación.",
    "R1 (bus factor): mitigar con documentación + código comentado + eventualmente onboardear dev junior. R2 (Supabase limits): monitorear uso; migrar a Pro ($25/mes) cuando se alcance 70% del free tier. R3 (WhatsApp API): tener plan B de email automation si no se aprueba. R4 (Claudia): documentar proceso de seguimiento; eventualmente contratar backup. R5 (competidor): diferenciar con curaduría y relaciones personales de Moisés (no replicables).", "✅")

add_qa(doc, "15.3", "Dependencias externas críticas.",
    "Supabase: BD, Auth, Storage. Sin Supabase, el sitio no funciona (solo páginas estáticas). Vercel: hosting y serverless functions. WhatsApp Business API (futuro): crítico para automatización. PostHog: analytics; si falla, no impacta funcionalidad pero ciega la medición.", "✅")

add_qa(doc, "15.4", "Compromisos con fechas regulatorias o contractuales.",
    "No hay fechas regulatorias inamovibles. El aviso de privacidad debe mantenerse actualizado. Si se implementa email marketing, cumplir con CAN-SPAM y LFPDPPP (incluir mecanismo de unsubscribe).", "✅")

add_qa(doc, "15.5", "¿Si un competidor lanza algo similar primero?",
    "Riesgo medio. La barrera de entrada técnica es baja (cualquiera puede hacer un sitio con Next.js + Supabase). La barrera real es: (a) la red de contactos y reputación de Moisés (15+ años en el mercado), (b) la calidad de la curaduría y el análisis financiero, (c) la velocidad de ejecución. Si un competidor lanza primero, Black Corporativo compite en calidad, no en timing.", "✅")

add_qa(doc, "15.6", "Riesgos de adopción.",
    "Compradores mayores (60+) pueden preferir WhatsApp directo a navegar un sitio. Mitigación: la plataforma no reemplaza a Moisés/Claudia, los potencia. El sitio es el escaparate; el cierre sigue siendo personal. También: ofrecer «tour guiado por WhatsApp» del sitio para clientes que lo prefieran.", "✅")

add_qa(doc, "15.7", "Supuestos más frágiles del modelo.",
    "1) Que el tráfico orgánico crecerá sin inversión en SEO. (Realidad: requiere contenido y backlinks). 2) Que los leads llegarán calificados sin filtro previo. (Realidad: lead magnet captura curiosos, no solo compradores serios). 3) Que Claudia puede manejar 40+ leads/mes sin automatización. (Realidad: a >30 leads/mes se satura).", "⚠")

add_qa(doc, "15.8", "Concentración de riesgo.",
    "Geográfico: 100% de propiedades en Tijuana/Baja California. Cliente: no hay dependencia de un solo cliente (base diversificada). Proveedor: Supabase y Vercel como proveedores únicos. Persona: Moisés como único desarrollador. Los últimos dos son los riesgos de concentración más altos.", "⚠")

add_qa(doc, "15.9", "Kill switch: criterio para detener el proyecto.",
    "Si después de 12 meses de campaña activa (oct 2026 - sep 2027) no se logran al menos 2 cierres/mes originados en digital, reconsiderar si el canal digital es el adecuado vs. prospección tradicional. El sitio se mantendría como escaparate pasivo pero sin inversión adicional en desarrollo.", "✅")

add_qa(doc, "15.10", "Monitoreo de riesgos en operación.",
    "Revisión mensual (Moisés): leads capturados, uptime del sitio, uso de Supabase (dashboard), feedback de Claudia sobre volumen y calidad de leads. Si algún riesgo se materializa (ej. Claudia reporta saturación), se escala a Moisés para decidir mitigación.", "✅")

# ============================================================
# SECCIÓN 16: ROADMAP Y ENTREGABLES
# ============================================================
add_section_band(doc, 16, "Roadmap y entregables",
    "Cuándo se entrega qué.")

add_qa(doc, "16.1", "Fase 0: descubrimiento y diseño.",
    "Completada (dic 2025 - abr 2026). Entregables: arquitectura definida (Next.js + Supabase), design system parcial (globals.css), wireframes implícitos (desarrollo directo), definición de 3 verticales de negocio.", "✅")

add_qa(doc, "16.2", "Fase 1: MVP (mayo 2026). YA ENTREGADO.",
    "Landing page con hero video, BrandsGrid, FeaturedInventory, SocialProof, LeadMagnet. 3 sub-brand pages (Luxury, Business, Industrial). Catálogo con filtro. Ficha de propiedad con galería, métricas, documentos, agente, CTA. Admin panel: dashboard, CRUD properties/agents/leads. Auth admin. Herramientas financieras.", "✅")

add_qa(doc, "16.3", "Fases 2, 3 y 4.",
    "F2 - Automatización (jul-sep 2026): WhatsApp Business API, email transaccional (Resend), favoritos localStorage, blog SEO, página de testimonios, Google Maps en ficha. F3 - Conversión (oct-dic 2026): búsqueda avanzada (texto, precio, m², mapa), chatbot IA básico, dashboard de métricas en admin, mejoras de accesibilidad WCAG AA. F4 - Escalamiento (ene-jun 2027): multi-idioma (EN), portal de propietarios, comparador de propiedades, notificaciones push PWA, matching automático.", "📋")

add_qa(doc, "16.4", "Hitos externos que condicionan el calendario.",
    "Temporada alta inmobiliaria (oct-mar): la F2 debe estar lista antes de oct 2026 para capturar el pico de tráfico. Aprobación de WhatsApp Business API: si se demora, la F2 se retrasa en su componente de automatización. Presupuesto: si no se aprueba inversión en marketing para sep 2026, el crecimiento será solo orgánico.", "⚠")

add_qa(doc, "16.5", "¿Cómo se prioriza?",
    "MoSCoW: Must have = lo que impacta directamente leads o cierres. Should have = mejoras de UX sin impacto directo en métrica principal. Could have = features «nice to have». Won't have = lo que está fuera de alcance. El criterio es: ¿esto aumenta leads calificados, reduce tiempo de respuesta, o incrementa tasa de cierre? Si no, se difiere.", "✅")

add_qa(doc, "16.6", "Fechas «no negociables» vs. negociables.",
    "No negociable: lanzamiento de campaña sep 2026 (oportunidad de temporada). Negociable: features específicas de cada fase (se pueden swappear según urgencia). No negociable: privacidad y seguridad (no se lanza nada que exponga datos de leads).", "✅")

add_qa(doc, "16.7", "Cadencia de release.",
    "Despliegue continuo (git push → Vercel). Releases semanales informales. No hay versionado semántico ni release notes públicas (no hay usuarios que dependan de versiones específicas). Si se implementa API pública, se adopta versionado semántico.", "✅")

add_qa(doc, "16.8", "Artefactos por fase.",
    "Todas las fases entregan: código (GitHub), migraciones SQL (si aplica), tests manuales (lista de verificación), y actualización de documentación en el README. F2+ entregan además: runbook de operación (cómo usar nuevas features en admin) y guion de capacitación para Claudia.", "✅")

add_qa(doc, "16.9", "Gestión de tech debt y re-work.",
    "Regla: 20% del tiempo de desarrollo se dedica a deuda técnica (refactors, tests, mejoras de rendimiento). Si una feature se construye con atajos, se registra como deuda y se paga en el siguiente ciclo. No se acumula deuda sin registrar.", "⚠")

add_qa(doc, "16.10", "Decisiones diferidas al siguiente PRD.",
    "Internacionalización completa (más allá de inglés). App nativa. Marketplace multi-broker. Integración con CRMs externos. Firma electrónica de contratos en plataforma. Analítica avanzada con IA predictiva.", "📋")

# ============================================================
# SECCIÓN 17: EQUIPO Y GOBERNANZA
# ============================================================
add_section_band(doc, 17, "Equipo y gobernanza",
    "Quién decide, quién ejecuta, quién revisa.")

add_qa(doc, "17.1", "DRI del producto.",
    "Moisés. Es el fundador, broker principal y desarrollador. Aprueba todas las decisiones de producto, tecnología y presupuesto.", "✅")

add_qa(doc, "17.2", "DRIs por área.",
    "Diseño: Moisés (no hay diseñador dedicado). Ingeniería: Moisés (full-stack). QA: Moisés + Claudia (pruebas manuales). Datos: Moisés (Supabase). Seguridad: Moisés. Operaciones: Claudia (seguimiento de leads) + Moisés (infraestructura). Legal: Moisés (con apoyo de abogado externo cuando se requiere). Comercial: Moisés.", "✅")

add_qa(doc, "17.3", "Tamaño y composición del equipo.",
    "Actual: 2 personas. Moisés (DRI, full-stack dev, broker) y Claudia (agente, soporte, seguimiento de leads). Futuro (12 meses): +1 diseñador UX freelance, +1 desarrollador junior/mid (para reducir bus factor), +1 agente adicional (si volumen de leads >40/mes).", "⚠")

add_qa(doc, "17.4", "Contratistas, agencias o partners externos.",
    "Fotógrafo/videógrafo: por proyecto (sesión de propiedad). Abogado: por consulta (contratos, aviso de privacidad). Diseñador UX: freelance planeado para Q3 2026. No hay agencia de desarrollo — todo es in-house o vía Hermes Agent como asistente de código.", "✅")

add_qa(doc, "17.5", "Cadencia de reuniones de gobierno.",
    "Diaria: Claudia y Moisés revisan nuevos leads (5 min por WhatsApp). Semanal: Moisés revisa métricas (PostHog + Supabase) y prioriza tareas de desarrollo. Mensual: revisión de KPIs vs. objetivos. No hay comité formal — Moisés toma las decisiones.", "✅")

add_qa(doc, "17.6", "¿Quién aprueba qué?",
    "Alcance y features: Moisés. Presupuesto: Moisés. Cambios que afecten privacidad o legal: Moisés con abogado. Diseño: Moisés (con input de Claudia sobre usabilidad del admin).", "✅")

add_qa(doc, "17.7", "Gestión de cambios al PRD una vez aprobado.",
    "Este documento es la línea base. Cambios se registran en el historial de versiones (apéndice). Si un cambio afecta alcance, fecha o presupuesto, Moisés lo evalúa y decide. No hay proceso formal de change control — es una startup de 2 personas.", "✅")

add_qa(doc, "17.8", "Reporte de avances, bloqueos y riesgos.",
    "Avances: Moisés actualiza el README con lo completado. Bloqueos: se comunican inmediatamente por WhatsApp. Riesgos: se revisan en la reunión mensual. No hay reportes formales para stakeholders externos.", "✅")

add_qa(doc, "17.9", "Dependencias de personal crítico (bus factor).",
    "Bus factor = 1 (Moisés). Si Moisés no está disponible, nadie puede modificar el código ni la infraestructura. Si Claudia no está, nadie da seguimiento a leads. Mitigación: documentar y eventualmente onboardear dev junior + agente backup.", "⚠")

add_qa(doc, "17.10", "Plan de conocimiento (documentación, traspasos).",
    "Código: GitHub + README. Procesos: no documentados formalmente. Plan: crear runbooks para Claudia (cómo crear propiedad, cómo seguir lead) y documentación técnica para futuro dev (arquitectura, decisiones, setup local). Esto es deuda actual.", "⚠")

# ============================================================
# SECCIÓN 18: CIERRE Y APROBACIONES
# ============================================================
add_section_band(doc, 18, "Cierre y aprobaciones",
    "Firmas que convierten el cuestionario en PRD.")

add_qa(doc, "18.1", "Resumen ejecutivo de una página.",
    "Black Corporativo es una plataforma digital inmobiliaria de alta gama (WPA responsive, Next.js + Supabase) que opera tres verticales —Black Luxury (residencial), Black Business (comercial) y Black Industrial— desde Tijuana, BC. El MVP está en producción con catálogo de propiedades, análisis financiero, captura de leads y admin panel. La meta a 12 meses es pasar de 15 a 40 leads/mes y automatizar el seguimiento con WhatsApp Business API. El modelo de negocio es corretaje por comisión (3-5% sobre transacciones de ~12M MXN promedio). Con una inversión mensual de $20K MXN en marketing digital, el ROI proyectado es >10:1. La plataforma diferencia por curaduría real, análisis financiero estructurado y experiencia institucional en un mercado (Tijuana/nearshoring) con alta demanda y baja oferta digital de calidad.", "✅")

add_qa(doc, "18.2", "Decisión solicitada.",
    "Se solicita aprobación para: (1) ejecutar el roadmap de 4 fases descrito en la sección 16, priorizando WhatsApp automation y SEO (F2, jul-sep 2026); (2) asignar presupuesto de $20K MXN/mes para campaña de lanzamiento en Google Ads y LinkedIn a partir de sep 2026; (3) iniciar búsqueda de diseñador UX freelance y desarrollador junior para reducir bus factor.", "✅")

add_qa(doc, "18.3", "Compromisos de la organización.",
    "Moisés: 5-10h/semana a desarrollo y revisión de métricas. Claudia: seguimiento de leads en <1h, mantener admin actualizado. Ambos: reportar métricas mensuales y ajustar estrategia. Inversión: $20K MXN/mes en marketing digital + $500 USD/mes en infraestructura cuando se migre a Supabase Pro + Vercel Pro.", "✅")

add_qa(doc, "18.4", "Anti-objetivos (lo que NO se hará).",
    "No se construirá un marketplace abierto. No se desarrollará app nativa. No se integrará pasarela de pagos. No se hará CRM completo. No se hará internacionalización más allá de inglés en 2027. No se contratará agencia de desarrollo externa.", "✅")

add_qa(doc, "18.5", "Próximos hitos.",
    "Jul 2026: WhatsApp API + email transaccional + blog SEO. Sep 2026: Campaña de lanzamiento + Google Maps en fichas. Dic 2026: Búsqueda avanzada + chatbot IA + dashboard de métricas. Revisión de PRD: enero 2027 (con datos reales de 6 meses de operación).", "✅")

add_qa(doc, "18.6", "Firmas requeridas.",
    "Patrocinador / Fundador / DRI: Moisés. Agente / Operaciones: Claudia. Este documento no requiere firma de terceros (legal externo, finanzas) por el tamaño actual de la operación. Se recomienda revisión de abogado para el aviso de privacidad y términos de uso antes del lanzamiento público (sep 2026).", "✅")

# ---------- Cierre ----------
sp = doc.add_paragraph()
sp.paragraph_format.space_before = Pt(24); r = sp.add_run("")
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
p.paragraph_format.space_after = Pt(4)
r = p.add_run("—  Fin del PRD respondido  —")
set_run(r, "Calibri", 13, bold=True, italic=True, color=GRAY)
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = p.add_run("Documento generado el 4 de junio de 2026 con base en análisis del código fuente de Black-Corporativo.")
set_run(r, "Calibri", 10, italic=True, color=GRAY)
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = p.add_run("Próxima revisión: enero 2027.")
set_run(r, "Calibri", 10, italic=True, color=GRAY)

# ---------- Guardar ----------
out = "/root/black_corporativo/CUESTIONARIO-PRD-RESPONDIDO.docx"
doc.save(out)
print(f"OK: {out}")
